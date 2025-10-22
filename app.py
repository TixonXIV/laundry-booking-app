from flask import Flask, render_template, request, redirect, url_for, session, flash, send_file
import sqlite3
from datetime import datetime
from docx import Document
from io import BytesIO

app = Flask(__name__)
app.secret_key = '9f3fd7d71cde58f23207fb1a98a0b266'

# Глобальное определение time_sort_key
def time_sort_key(ts):
    try:
        start_time = ts.split('-')[0]
        h, m = map(int, start_time.split(':'))
        return h * 60 + m
    except:
        return 0

# Регистрация функции time_sort_key для SQLite
def register_sql_functions(conn):
    conn.create_function("time_sort_key", 1, time_sort_key)

# Инициализация базы данных
def init_db():
    conn = sqlite3.connect('laundry.db')
    register_sql_functions(conn)
    c = conn.cursor()
    c.execute('''CREATE TABLE IF NOT EXISTS users
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  surname TEXT NOT NULL,
                  room TEXT NOT NULL,
                  UNIQUE(surname, room))''')
    c.execute('''CREATE TABLE IF NOT EXISTS days
                 (name TEXT PRIMARY KEY,
                  order_num INTEGER UNIQUE)''')
    c.execute('''CREATE TABLE IF NOT EXISTS machines
                 (number INTEGER PRIMARY KEY,
                  status TEXT DEFAULT 'active')''')
    c.execute('''CREATE TABLE IF NOT EXISTS slots
                 (id INTEGER PRIMARY KEY AUTOINCREMENT,
                  day_name TEXT NOT NULL,
                  time_slot TEXT NOT NULL,
                  machine_number INTEGER NOT NULL,
                  user_id INTEGER,
                  FOREIGN KEY(day_name) REFERENCES days(name),
                  FOREIGN KEY(machine_number) REFERENCES machines(number),
                  UNIQUE(day_name, time_slot, machine_number))''')
    conn.commit()
    conn.close()

init_db()

# Стандартные временные слоты
default_time_slots = [
    '7:00-9:00',
    '10:00-12:00',
    '13:00-15:00',
    '16:00-18:00',
    '19:00-21:00',
    '22:00-24:00'
]

def get_days():
    conn = sqlite3.connect('laundry.db')
    register_sql_functions(conn)
    c = conn.cursor()
    c.execute('SELECT name FROM days ORDER BY order_num')
    days = [row[0] for row in c.fetchall()]
    conn.close()
    return days

def get_machines():
    conn = sqlite3.connect('laundry.db')
    register_sql_functions(conn)
    c = conn.cursor()
    c.execute('SELECT number, status FROM machines ORDER BY number')
    machines = [(row[0], row[1]) for row in c.fetchall()]
    conn.close()
    return machines

def get_slots(for_admin=False):
    conn = sqlite3.connect('laundry.db')
    register_sql_functions(conn)
    c = conn.cursor()
    if for_admin:
        c.execute('SELECT day_name, time_slot, machine_number, user_id, m.status FROM slots s LEFT JOIN machines m ON s.machine_number = m.number ORDER BY day_name, time_sort_key(time_slot), machine_number')
    else:
        c.execute('SELECT day_name, time_slot, machine_number, user_id, m.status FROM slots s LEFT JOIN machines m ON s.machine_number = m.number WHERE m.status = "active" ORDER BY day_name, time_sort_key(time_slot), machine_number')
    slots = c.fetchall()
    conn.close()
    return slots

def is_user_booked(user_id):
    conn = sqlite3.connect('laundry.db')
    register_sql_functions(conn)
    c = conn.cursor()
    c.execute('SELECT * FROM slots WHERE user_id = ?', (user_id,))
    booked = c.fetchone()
    conn.close()
    return booked is not None

@app.route('/', methods=['GET', 'POST'])
def register():
    if request.method == 'POST':
        surname = request.form['surname']
        room = request.form['room']
        conn = sqlite3.connect('laundry.db')
        register_sql_functions(conn)
        c = conn.cursor()
        try:
            c.execute('INSERT INTO users (surname, room) VALUES (?, ?)', (surname, room))
            conn.commit()
            user_id = c.lastrowid
            session['user_id'] = user_id
            session['surname'] = surname
            session['room'] = room
            return redirect(url_for('schedule'))
        except sqlite3.IntegrityError:
            c.execute('SELECT id FROM users WHERE surname = ? AND room = ?', (surname, room))
            user = c.fetchone()
            if user:
                session['user_id'] = user[0]
                session['surname'] = surname
                session['room'] = room
                return redirect(url_for('schedule'))
            else:
                flash('Ошибка регистрации')
        conn.close()
    return render_template('register.html')

@app.route('/schedule', methods=['GET', 'POST'])
def schedule():
    if 'user_id' not in session:
        return redirect(url_for('register'))
    if is_user_booked(session['user_id']):
        flash('Вы уже записаны. Только одна запись на пользователя.')
    
    days = get_days()
    machines = get_machines()
    active_machines = [m for m, s in machines if s == 'active']
    num_machines = len(active_machines)
    
    slots = get_slots(for_admin=False)
    current_time_slots = set(ts for _, ts, _, _, _ in slots)
    time_slots = sorted(list(current_time_slots), key=time_sort_key)
    
    schedule_data = {day: {ts: [None] * num_machines for ts in time_slots} for day in days}
    machine_map = {active_machines[i]: i for i in range(num_machines)}
    
    for day, ts, machine, user_id, status in slots:
        if day in schedule_data and ts in schedule_data[day]:
            idx = machine_map.get(machine)
            if idx is not None:
                if user_id:
                    conn = sqlite3.connect('laundry.db')
                    register_sql_functions(conn)
                    c = conn.cursor()
                    c.execute('SELECT surname, room FROM users WHERE id = ?', (user_id,))
                    user = c.fetchone()
                    conn.close()
                    schedule_data[day][ts][idx] = f"{user[0]} {user[1]}"
                else:
                    schedule_data[day][ts][idx] = None
    
    if request.method == 'POST':
        if is_user_booked(session['user_id']):
            return redirect(url_for('schedule'))
        day = request.form['day']
        ts = request.form['time_slot']
        machine = int(request.form['machine'])
        conn = sqlite3.connect('laundry.db')
        register_sql_functions(conn)
        c = conn.cursor()
        c.execute('''SELECT s.user_id, m.status FROM slots s 
                     JOIN machines m ON s.machine_number = m.number 
                     WHERE s.day_name = ? AND s.time_slot = ? AND s.machine_number = ?''', (day, ts, machine))
        slot = c.fetchone()
        if slot and slot[0] is None and slot[1] == 'active':
            c.execute('UPDATE slots SET user_id = ? WHERE day_name = ? AND time_slot = ? AND machine_number = ?', (session['user_id'], day, ts, machine))
            conn.commit()
            flash('Запись успешна!')
        else:
            flash('Слот занят, не существует или машина отключена.')
        conn.close()
        return redirect(url_for('schedule'))
    
    return render_template('schedule.html', schedule_data=schedule_data, days=days, time_slots=time_slots, machines=active_machines)

@app.route('/admin_login', methods=['GET', 'POST'])
def admin_login():
    if session.get('admin'):
        return redirect(url_for('admin'))
    if request.method == 'POST':
        username = request.form['username']
        password = request.form['password']
        if username == 'admin' and password == 'admin123':  # Измените на безопасный пароль
            session['admin'] = True
            return redirect(url_for('admin'))
        else:
            flash('Неверный логин или пароль')
    return render_template('admin_login.html')

@app.route('/admin_logout')
def admin_logout():
    session.pop('admin', None)
    return redirect(url_for('admin_login'))

@app.route('/admin', methods=['GET', 'POST'])
def admin():
    if not session.get('admin'):
        return redirect(url_for('admin_login'))
    
    days = get_days()
    machines = get_machines()
    num_machines = len(machines)
    
    slots = get_slots(for_admin=True)
    current_time_slots = set(ts for _, ts, _, _, _ in slots)
    time_slots = sorted(list(current_time_slots), key=time_sort_key)
    
    schedule_data = {day: {ts: [None] * num_machines for ts in time_slots} for day in days}
    machine_map = {machines[i][0]: i for i in range(num_machines)}
    
    for day, ts, machine, user_id, status in slots:
        if day in schedule_data and ts in schedule_data[day]:
            idx = machine_map.get(machine)
            if idx is not None:
                if status == 'disabled':
                    schedule_data[day][ts][idx] = 'Отключена'
                elif user_id:
                    conn = sqlite3.connect('laundry.db')
                    register_sql_functions(conn)
                    c = conn.cursor()
                    c.execute('SELECT surname, room FROM users WHERE id = ?', (user_id,))
                    user = c.fetchone()
                    conn.close()
                    schedule_data[day][ts][idx] = f"{user[0]} {user[1]}"
                else:
                    schedule_data[day][ts][idx] = None
    
    if request.method == 'POST':
        action = request.form.get('action')
        conn = sqlite3.connect('laundry.db')
        register_sql_functions(conn)
        c = conn.cursor()
        
        if action == 'reset':
            c.execute('UPDATE slots SET user_id = NULL')
            conn.commit()
            flash('Таблица сброшена.')
        
        elif action == 'factory_reset':
            # Сброс к заводским настройкам: сохраняем users, days, machines, восстанавливаем slots с базовым временем
            c.execute('DELETE FROM slots')
            conn.commit()
            days = get_days()
            machines = [m for m, _ in get_machines()]
            for day in days:
                for ts in default_time_slots:
                    for machine in machines:
                        c.execute('INSERT INTO slots (day_name, time_slot, machine_number, user_id) VALUES (?, ?, ?, NULL)', (day, ts, machine))
            conn.commit()
            flash('Таблица возвращена к заводским настройкам.')
        
        elif action == 'add_time':
            new_time = request.form['new_time']
            if new_time not in current_time_slots:
                for day in days:
                    for machine, _ in machines:
                        c.execute('INSERT INTO slots (day_name, time_slot, machine_number, user_id) VALUES (?, ?, ?, NULL)', (day, new_time, machine))
                conn.commit()
                flash('Новый временной слот добавлен.')
        
        elif action == 'remove_time':
            remove_time = request.form['remove_time']
            if remove_time in current_time_slots:
                c.execute('DELETE FROM slots WHERE time_slot = ?', (remove_time,))
                conn.commit()
                flash('Временной слот удален.')
        
        elif action == 'edit_time':
            old_time = request.form['old_time']
            new_time = request.form['new_time']
            if old_time in current_time_slots:
                c.execute('UPDATE slots SET time_slot = ? WHERE time_slot = ?', (new_time, old_time))
                conn.commit()
                flash('Временной слот изменен.')
        
        elif action == 'add_day':
            new_day = request.form['new_day']
            order_num = len(days) + 1
            try:
                c.execute('INSERT INTO days (name, order_num) VALUES (?, ?)', (new_day, order_num))
                for ts in time_slots:
                    for machine, _ in machines:
                        c.execute('INSERT INTO slots (day_name, time_slot, machine_number, user_id) VALUES (?, ?, ?, NULL)', (new_day, ts, machine))
                conn.commit()
                flash('Новый день добавлен.')
            except sqlite3.IntegrityError:
                flash('День уже существует.')
        
        elif action == 'edit_day':
            old_day = request.form['old_day']
            new_day = request.form['new_day']
            c.execute('UPDATE days SET name = ? WHERE name = ?', (new_day, old_day))
            c.execute('UPDATE slots SET day_name = ? WHERE day_name = ?', (new_day, old_day))
            conn.commit()
            flash('День изменен.')
        
        elif action == 'remove_day':
            remove_day = request.form['remove_day']
            c.execute('DELETE FROM slots WHERE day_name = ?', (remove_day,))
            c.execute('DELETE FROM days WHERE name = ?', (remove_day,))
            conn.commit()
            flash('День удален.')
        
        elif action == 'add_machine':
            new_number = max([m for m, _ in machines], default=0) + 1
            c.execute('INSERT INTO machines (number, status) VALUES (?, "active")', (new_number,))
            for day in days:
                for ts in time_slots:
                    c.execute('INSERT INTO slots (day_name, time_slot, machine_number, user_id) VALUES (?, ?, ?, NULL)', (day, ts, new_number))
            conn.commit()
            flash('Новая машина добавлена.')
        
        elif action == 'remove_machine':
            remove_machine = int(request.form['remove_machine'])
            c.execute('DELETE FROM slots WHERE machine_number = ?', (remove_machine,))
            c.execute('DELETE FROM machines WHERE number = ?', (remove_machine,))
            conn.commit()
            flash('Машина удалена.')
        
        elif action == 'toggle_machine':
            machine_num = int(request.form['machine_num'])
            c.execute('SELECT status FROM machines WHERE number = ?', (machine_num,))
            current = c.fetchone()
            if current:
                new_status = 'disabled' if current[0] == 'active' else 'active'
                c.execute('UPDATE machines SET status = ? WHERE number = ?', (new_status, machine_num))
                if new_status == 'disabled':
                    c.execute('UPDATE slots SET user_id = NULL WHERE machine_number = ? AND user_id IS NOT NULL', (machine_num,))
                conn.commit()
                flash(f'Машина {machine_num} {"отключена" if new_status == "disabled" else "включена"}.')
        
        elif action == 'edit':
            day = request.form['day']
            ts = request.form['time_slot']
            machine = int(request.form['machine'])
            new_value = request.form['new_value'].strip()
            
            c.execute('SELECT status FROM machines WHERE number = ?', (machine,))
            status = c.fetchone()
            if status and status[0] == 'disabled':
                flash('Машина отключена, нельзя редактировать.')
            else:
                if new_value == '':
                    c.execute('UPDATE slots SET user_id = NULL WHERE day_name = ? AND time_slot = ? AND machine_number = ?', (day, ts, machine))
                    conn.commit()
                    flash('Запись пользователя удалена.')
                else:
                    try:
                        surname, room = new_value.split(' ')
                        c.execute('SELECT id FROM users WHERE surname = ? AND room = ?', (surname, room))
                        user = c.fetchone()
                        if not user:
                            c.execute('INSERT INTO users (surname, room) VALUES (?, ?)', (surname, room))
                            conn.commit()
                            c.execute('SELECT id FROM users WHERE surname = ? AND room = ?', (surname, room))
                            user = c.fetchone()
                        if user:
                            c.execute('UPDATE slots SET user_id = ? WHERE day_name = ? AND time_slot = ? AND machine_number = ?', (user[0], day, ts, machine))
                            conn.commit()
                            flash('Слот отредактирован.')
                        else:
                            flash('Не удалось создать пользователя.')
                    except ValueError:
                        flash('Неверный формат: Фамилия Комната')
        
        elif action == 'export_word':
            days = get_days()
            machines = get_machines()
            num_machines = len(machines)
            slots = get_slots(for_admin=True)
            time_slots_set = set(ts for _, ts, _, _, _ in slots)
            time_slots = sorted(time_slots_set, key=time_sort_key)
            
            doc = Document()
            doc.add_heading('Расписание стирки', 0)
            
            for day in days:
                doc.add_heading(day, 1)
                table = doc.add_table(rows=1, cols=num_machines + 1)
                table.style = 'Table Grid'
                
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = 'Время'
                for i, (m_num, _) in enumerate(machines):
                    hdr_cells[i + 1].text = f'Машина {m_num}'
                
                for ts in time_slots:
                    row_cells = table.add_row().cells
                    row_cells[0].text = ts
                    for i, (m_num, m_status) in enumerate(machines):
                        user_str = ''
                        for slot_day, slot_ts, slot_m, user_id, status in slots:
                            if slot_day == day and slot_ts == ts and slot_m == m_num:
                                if status == 'disabled':
                                    user_str = 'Отключена'
                                elif user_id:
                                    conn2 = sqlite3.connect('laundry.db')
                                    register_sql_functions(conn2)
                                    c2 = conn2.cursor()
                                    c2.execute('SELECT surname, room FROM users WHERE id = ?', (user_id,))
                                    user = c2.fetchone()
                                    conn2.close()
                                    if user:
                                        user_str = f"{user[0]} {user[1]}"
                                break
                        row_cells[i + 1].text = user_str
            
            buf = BytesIO()
            doc.save(buf)
            buf.seek(0)
            return send_file(buf, as_attachment=True, download_name='schedule.docx', mimetype='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        
        conn.close()
        return redirect(url_for('admin'))
    
    return render_template('admin.html', schedule_data=schedule_data, days=days, time_slots=time_slots, machines=[m for m, _ in machines])

def init_data():
    conn = sqlite3.connect('laundry.db')
    register_sql_functions(conn)
    c = conn.cursor()
    
    c.execute('SELECT COUNT(*) FROM days')
    if c.fetchone()[0] == 0:
        initial_days = ['Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота', 'Воскресенье']
        for i, day in enumerate(initial_days, 1):
            c.execute('INSERT INTO days (name, order_num) VALUES (?, ?)', (day, i))
    
    c.execute('SELECT COUNT(*) FROM machines')
    if c.fetchone()[0] == 0:
        for i in range(1, 6):
            c.execute('INSERT INTO machines (number, status) VALUES (?, "active")', (i,))
    
    c.execute('SELECT COUNT(*) FROM slots')
    if c.fetchone()[0] == 0:
        days = get_days()
        machines = [m for m, _ in get_machines()]
        for day in days:
            for ts in default_time_slots:
                for machine in machines:
                    c.execute('INSERT INTO slots (day_name, time_slot, machine_number, user_id) VALUES (?, ?, ?, NULL)', (day, ts, machine))
    
    conn.commit()
    conn.close()

init_data()

if __name__ == '__main__':
    app.run(debug=True)